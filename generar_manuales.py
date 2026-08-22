import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text):
    title = doc.add_heading(text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(doc, text, level=1):
    doc.add_heading(text, level)

def add_paragraph(doc, text):
    doc.add_paragraph(text)

def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def add_image_if_exists(doc, image_path, width_inches=5.0):
    if os.path.exists(image_path):
        try:
            doc.add_picture(image_path, width=Inches(width_inches))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Error adding image {image_path}: {e}")

img_dir = r"C:\Users\cachi\.gemini\antigravity-ide\brain\da98eeb5-4723-4927-8b21-dcd820f13824\.user_uploaded"

def create_manual_tecnico():
    doc = Document()
    add_title(doc, "MANUAL TÉCNICO\nSistema de Monitoreo y Alerta Temprana para Riesgos Climáticos (ClimaSafe)")
    
    # CHAPTER 1
    add_heading(doc, "1. Introducción y Objetivos", level=1)
    add_paragraph(doc, "El presente manual técnico detalla la arquitectura, tecnologías y configuración del sistema ClimaSafe. Este sistema integral ha sido desarrollado meticulosamente para proveer monitoreo en tiempo real, alertas tempranas y registro histórico de variables climáticas en la región de Villa Canales, Guatemala. La aplicación busca dotar a los administradores y operadores de herramientas proactivas ante posibles desastres naturales, como inundaciones o sequías extremas.")
    add_paragraph(doc, "El alcance de este documento cubre el diseño de la arquitectura, requerimientos estrictos del servidor, modelado de la base de datos relacional y los comandos exactos para garantizar un despliegue exitoso en un entorno de producción (Nube).")
    doc.add_page_break()

    # CHAPTER 2
    add_heading(doc, "2. Requisitos del Sistema (Software y Hardware)", level=1)
    add_heading(doc, "2.1 Hardware (Servidor VPS de Producción)", level=2)
    add_paragraph(doc, "Para asegurar la disponibilidad y estabilidad del motor de base de datos SQL Server 2022 y los contenedores de la aplicación, se establecen los siguientes requisitos mínimos de hardware para el Droplet (DigitalOcean) o instancia EC2 (AWS):")
    add_bullet(doc, "Procesador (CPU): 1 vCPU (2.5 GHz o superior).")
    add_bullet(doc, "Memoria RAM: 2 GB (Requisito estricto, ya que el contenedor mssql-server se reiniciará por falta de memoria con menos de 2GB).")
    add_bullet(doc, "Almacenamiento: 25 GB SSD (Para acomodar volúmenes de Docker, logs del sistema y datos históricos de sensores).")
    add_bullet(doc, "Red: Dirección IP Pública IPv4 dedicada para mapeo de dominios.")

    add_heading(doc, "2.2 Software (Entorno de Servidor y Cliente)", level=2)
    add_paragraph(doc, "El sistema operativo host y las dependencias necesarias son:")
    add_bullet(doc, "Sistema Operativo Host: Ubuntu Linux 24.04 LTS (Recomendado).")
    add_bullet(doc, "Motor de Contenedores: Docker Engine (v24 o superior).")
    add_bullet(doc, "Orquestador Local: Docker Compose (v2).")
    add_bullet(doc, "Control de Versiones: Git CLI.")
    add_bullet(doc, "Navegador Cliente: Google Chrome, Safari, Microsoft Edge (Versiones recientes con soporte para WebSockets).")
    
    add_image_if_exists(doc, os.path.join(img_dir, "media_1787201389942.png"))
    doc.add_page_break()

    # CHAPTER 3
    add_heading(doc, "3. Tecnologías Utilizadas (Stack Tecnológico)", level=1)
    add_paragraph(doc, "La plataforma ha sido construida utilizando un esquema Cliente-Servidor Desacoplado, operando en contenedores Docker para garantizar la portabilidad.")
    
    add_heading(doc, "3.1 Frontend (Capa de Presentación)", level=2)
    add_paragraph(doc, "Desarrollado en Angular 17+ (Framework SPA de Google), utilizando TypeScript para el tipado estricto. La interfaz de usuario es 100% responsiva (Mobile-First) utilizando CSS3 moderno. La comunicación bidireccional en tiempo real para la actualización de las tarjetas de los sensores se realiza a través de la librería @microsoft/signalr.")
    
    add_heading(doc, "3.2 Backend (Capa Lógica y API)", level=2)
    add_paragraph(doc, "Construido con .NET 8, el framework de alto rendimiento de Microsoft. Se emplea ASP.NET Core para exponer una API RESTful. La arquitectura interna sigue patrones de inyección de dependencias y servicios asíncronos (HostedServices) para simular el hardware IoT y despachar notificaciones vía SignalR (WebSockets).")
    
    add_heading(doc, "3.3 Base de Datos y Almacenamiento", level=2)
    add_paragraph(doc, "Se emplea Microsoft SQL Server 2022 Developer Edition ejecutándose en un entorno Linux (mcr.microsoft.com/mssql/server:2022-latest). El almacenamiento persistente se garantiza a través de Docker Volumes (sql_data), lo que significa que aunque el contenedor se destruya, los datos permanecen intactos en el host.")
    
    add_image_if_exists(doc, os.path.join(img_dir, "media_1787030198567.png"))
    doc.add_page_break()

    # CHAPTER 4
    add_heading(doc, "4. Arquitectura de Despliegue y Orquestación (Docker)", level=1)
    add_paragraph(doc, "El archivo docker-compose.yml es el núcleo operativo de la infraestructura. Define cuatro servicios esenciales que conforman la red privada 'clima_network':")
    
    add_heading(doc, "4.1 Servicio: database (SQL Server)", level=2)
    add_paragraph(doc, "Imagen oficial de Microsoft. Expone el puerto 1433 y configura la contraseña SA mediante variables de entorno (MSSQL_SA_PASSWORD).")
    
    add_heading(doc, "4.2 Servicio: db-init (Seed Script)", level=2)
    add_paragraph(doc, "Es un contenedor efímero. Su único propósito es esperar a que SQL Server esté listo, conectarse utilizando sqlcmd, y ejecutar el archivo schema.sql que contiene la creación de las tablas, los procedimientos almacenados y los usuarios iniciales (Semilla).")
    
    add_heading(doc, "4.3 Servicio: backend (API)", level=2)
    add_paragraph(doc, "Se construye mediante un Dockerfile de múltiples etapas (Multi-stage build). Utiliza el SDK de .NET para compilar y publica en la imagen ligera de ASP.NET runtime, exponiendo el puerto 5000 para el tráfico HTTP.")
    
    add_heading(doc, "4.4 Servicio: frontend (Nginx)", level=2)
    add_paragraph(doc, "La aplicación Angular se compila (ng build) dentro del contenedor Node.js y los archivos estáticos resultantes se sirven a través del servidor web ultraligero Nginx, el cual expone el puerto 80 hacia el mundo exterior (Internet).")
    
    add_image_if_exists(doc, os.path.join(img_dir, "media_1787027425900.png"))
    doc.add_page_break()

    # CHAPTER 5
    add_heading(doc, "5. Diseño de Base de Datos y Diccionario de Datos", level=1)
    add_paragraph(doc, "El esquema ClimaDB está diseñado para ser altamente relacional, normalizado (hasta 3FN) y optimizado para inserciones a alta velocidad (provenientes de los sensores IoT).")
    
    add_image_if_exists(doc, os.path.join(img_dir, "media_1786840716711.png"))

    add_heading(doc, "5.1 Tabla: Usuarios", level=2)
    add_paragraph(doc, "Guarda la identidad de los operadores y administradores. Campos principales: Nombre, Email, PasswordHash (Almacenado como HASH SHA2-256), Rol, Activo.")
    
    add_heading(doc, "5.2 Tabla: Sensores", level=2)
    add_paragraph(doc, "Entidad física desplegada en Villa Canales. Campos: UsuarioId (FK), TipoSensor, UnidadMedida, ValorMinimo, ValorMaximo, Latitud, Longitud.")
    
    add_heading(doc, "5.3 Tabla: Lecturas", level=2)
    add_paragraph(doc, "Registra cada pulso del sensor. Campos: SensorId (FK), Valor, FechaHora. (Esta tabla crece rápidamente).")
    
    add_heading(doc, "5.4 Tabla: Alertas", level=2)
    add_paragraph(doc, "Registra eventos fuera de los límites. Campos: SensorId, Mensaje, Nivel (Crítico, Medio).")
    
    add_heading(doc, "5.5 Tablas: HistorialEventos y BitacoraAcciones", level=2)
    add_paragraph(doc, "Mantiene la trazabilidad y auditoría de accesos al sistema.")
    doc.add_page_break()

    # CHAPTER 6
    add_heading(doc, "6. Seguridad y Prevención de Riesgos Informáticos", level=1)
    add_paragraph(doc, "El sistema fue diseñado considerando las directrices OWASP Top 10:")
    add_bullet(doc, "Inmunidad a Inyección SQL: Absolutamente todas las consultas (INSERT, SELECT) se realizan mediante Stored Procedures con parámetros tipados (@Parametro). No existe concatenación de strings en el Backend.")
    add_bullet(doc, "Autenticación Criptográfica: Las contraseñas nunca viajan ni se almacenan en texto claro. Se utiliza HASHBYTES('SHA2_256', Password) con conversión binaria.")
    add_bullet(doc, "Exposición de Puertos: Solo los puertos 80 (Web), 5000 (API) y 1433 (DB remota) están accesibles, protegiendo las redes internas de Docker.")
    
    add_image_if_exists(doc, os.path.join(img_dir, "media_1787203723011.png"))
    doc.add_page_break()

    # CHAPTER 7
    add_heading(doc, "7. Guía de Despliegue en Producción (DigitalOcean / DuckDNS)", level=1)
    add_paragraph(doc, "Procedimiento estricto para poner el sistema en vivo en la nube pública:")
    add_paragraph(doc, "Paso 1: Configurar el servidor DNS dinámico (DuckDNS). Asignar la IP del VPS (ej. 198.199.65.77) al subdominio (ej. villacanales-clima.duckdns.org).")
    add_paragraph(doc, "Paso 2: Acceder al VPS mediante protocolo SSH (Secure Shell).")
    add_bullet(doc, "ssh root@198.199.65.77")
    add_paragraph(doc, "Paso 3: Instalar Git y Docker Compose en la máquina anfitriona.")
    add_paragraph(doc, "Paso 4: Clonar el repositorio y acceder a la carpeta de despliegue:")
    add_bullet(doc, "git clone https://github.com/maxair180/alerta-temprana-clima.git")
    add_bullet(doc, "cd alerta-temprana-clima/docker")
    add_paragraph(doc, "Paso 5: Construir e iniciar el enjambre de microservicios:")
    add_bullet(doc, "docker compose up -d --build")
    add_paragraph(doc, "Paso 6: Monitorear el despliegue con 'docker ps'. Verificar que los 3 contenedores (frontend, backend, sqlserver) permanezcan en estado 'Up'.")
    add_paragraph(doc, "El sistema estará automáticamente disponible de forma mundial a través del nombre de dominio.")

    doc.save('Manual_Tecnico_ClimaSafe.docx')
    print("Manual_Tecnico_ClimaSafe.docx generado (Extendido).")

def create_manual_usuario():
    doc = Document()
    add_title(doc, "MANUAL DE USUARIO\nSistema de Monitoreo y Alerta Temprana para Riesgos Climáticos (ClimaSafe)")
    
    # CHAPTER 1
    add_heading(doc, "1. Introducción y Propósito del Sistema", level=1)
    add_paragraph(doc, "El Sistema ClimaSafe ha sido diseñado con una interfaz amigable (User-Friendly) orientada a proteger a la población de Villa Canales mediante el monitoreo tecnológico avanzado. Este manual le guiará a través de todas las funcionalidades, desde el inicio de sesión hasta la interpretación de mapas climatológicos y reportes de auditoría.")
    add_paragraph(doc, "Usted podrá anticiparse a lluvias torrenciales, crecidas de ríos y olas de calor extremas gracias a la automatización de nuestros sensores en tiempo real.")
    doc.add_page_break()

    # CHAPTER 2
    add_heading(doc, "2. Acceso al Sistema y Roles de Seguridad", level=1)
    add_paragraph(doc, "Para garantizar la seguridad de los datos climáticos, el sistema requiere autenticación. Para ingresar:")
    add_bullet(doc, "1. Abra el navegador de su dispositivo móvil o computadora de escritorio.")
    add_bullet(doc, "2. Escriba la dirección web: http://villacanales-clima.duckdns.org")
    add_bullet(doc, "3. Se desplegará la pantalla de Login con protección Anti-Bots.")
    
    add_image_if_exists(doc, os.path.join(img_dir, "media_1787030198567.png"))
    
    add_heading(doc, "2.1 Roles Disponibles", level=2)
    add_paragraph(doc, "El Administrador tiene acceso total al sistema, incluyendo las bitácoras del servidor. El Operador tiene acceso limitado, enfocado estrictamente en monitorear los sensores y reaccionar a las alertas.")
    add_paragraph(doc, "Credenciales de Ejemplo para pruebas Universitarias:")
    add_bullet(doc, "Admin: ccachinm@miumg.edu.gt (Clave: Admin2026!)")
    add_bullet(doc, "Operador: cgarciaf11@miumg.edu.gt (Clave: YAYA@2026)")
    add_bullet(doc, "Operador: mlorenzanaa@miumg.edu.gt (Clave: ROSSE@2026)")
    doc.add_page_break()

    # CHAPTER 3
    add_heading(doc, "3. Navegación por el Dashboard Interactivo", level=1)
    add_paragraph(doc, "Una vez logueado, ingresará al corazón del sistema: El Dashboard. Esta pantalla está dividida estratégicamente en tres secciones vitales para la toma de decisiones rápidas.")
    
    add_heading(doc, "3.1 Tarjetas de Sensores en Tiempo Real", level=2)
    add_paragraph(doc, "En la parte superior, observará una cuadrícula con sensores simulados:")
    add_bullet(doc, "Termómetro (Temperatura en °C)")
    add_bullet(doc, "Pluviómetro (Nivel de Lluvia en mm)")
    add_bullet(doc, "Anemómetro (Viento en km/h)")
    add_bullet(doc, "Medidor de Río (Nivel del agua en metros)")
    add_paragraph(doc, "Estos valores cambian cada segundo de forma automática. Si un sensor detecta un valor peligroso (ejemplo: Río a 5 metros), la tarjeta empezará a parpadear en ROJO de forma intermitente para captar su atención inmediatamente.")
    
    add_image_if_exists(doc, os.path.join(img_dir, "media_1787027425900.png"))
    doc.add_page_break()

    # CHAPTER 4
    add_heading(doc, "4. Módulo Geográfico (Mapa de Sensores)", level=1)
    add_paragraph(doc, "Justo debajo de las tarjetas, se integra un mapa satelital de la región de Villa Canales. Los pines en el mapa representan las ubicaciones físicas de los sensores de campo (Aldea El Tablón, Boca del Monte, etc.).")
    add_paragraph(doc, "Esta herramienta permite a las autoridades coordinar esfuerzos de evacuación sabiendo exactamente qué comunidad está bajo amenaza climática en base a las coordenadas GPS de cada sensor.")
    
    add_image_if_exists(doc, os.path.join(img_dir, "media_1787201982719.png"))
    doc.add_page_break()

    # CHAPTER 5
    add_heading(doc, "5. Gestión de Notificaciones y Alertas", level=1)
    add_paragraph(doc, "El panel lateral (Sidebar) derecho funciona como un centro de notificaciones Push. Cada vez que el servidor detecta una anomalía, inyecta una notificación de alto contraste en este panel sin necesidad de que usted recargue la página.")
    add_bullet(doc, "Alerta Amarilla: Indica variaciones climáticas moderadas.")
    add_bullet(doc, "Alerta Roja (Crítica): Indica un riesgo inminente para la población civil (Ej. Huracán o Inundación). Requiere intervención inmediata.")
    
    add_image_if_exists(doc, os.path.join(img_dir, "media_1787202185496.png"))
    doc.add_page_break()

    # CHAPTER 6
    add_heading(doc, "6. Módulos de Auditoría (Bitácora e Historial)", level=1)
    add_paragraph(doc, "Para garantizar la transparencia gubernamental o corporativa, el sistema nunca borra datos, sino que los almacena para consulta histórica.")
    
    add_heading(doc, "6.1 Historial de Eventos Climáticos", level=2)
    add_paragraph(doc, "Un reporte tabular que lista absolutamente todas las alarmas que se han disparado en el último mes. Útil para que los analistas y meteorólogos calculen patrones climáticos en Villa Canales a lo largo del año.")
    
    add_heading(doc, "6.2 Bitácora de Acciones (Seguridad)", level=2)
    add_paragraph(doc, "Este submódulo es exclusivo para el Administrador. Muestra una tabla detallada de 'Quién hizo qué y a qué hora'.")
    add_bullet(doc, "Por ejemplo: Verificará si un Operador inició sesión correctamente, si hubo intentos fallidos, o si alguien realizó una modificación en los umbrales de los sensores.")
    
    add_image_if_exists(doc, os.path.join(img_dir, "media_1787364669059.png"))
    
    doc.add_page_break()
    add_heading(doc, "7. Cierre de Sesión Seguro", level=1)
    add_paragraph(doc, "Por motivos de privacidad, asegúrese de presionar el botón 'Cerrar Sesión' ubicado en la esquina superior derecha cada vez que se levante de su estación de trabajo. El sistema purgará sus credenciales de la memoria del navegador de inmediato.")

    doc.save('Manual_De_Usuario_ClimaSafe.docx')
    print("Manual_De_Usuario_ClimaSafe.docx generado (Extendido).")

if __name__ == '__main__':
    create_manual_tecnico()
    create_manual_usuario()
